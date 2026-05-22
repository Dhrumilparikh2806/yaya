from docx import Document as DocxDocument

from app.processors.base import BaseProcessor


class DOCXProcessor(BaseProcessor):
    def extract(self) -> str:
        doc = DocxDocument(self.job.file_path)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            parts.append(self._table_to_markdown(rows))
        return "\n\n".join(parts)

    def summarise(self, text: str, db) -> dict:
        if len(text) > 30000:
            self.log.warning("docx_text_truncated", original_len=len(text), truncated_to=30000)
            text = text[:30000]

        prompt = f"""You are a document analyst. Analyse the following document text and return ONLY valid JSON.
No preamble, no markdown code blocks, just raw JSON.

Return this exact structure:
{{
  "title": "document title or filename",
  "document_type": "report|contract|paper|manual|other",
  "summary": "2-3 sentence summary",
  "key_points": ["point 1", "point 2", ...],
  "risks": ["risk 1", ...],
  "sections": ["list of section headings found in the document"],
  "entities": ["company names, person names, product names mentioned"]
}}

Document text:
{text}
"""
        return self._call_gemini_json(prompt, db)
